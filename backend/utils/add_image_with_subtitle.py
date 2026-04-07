from typing import Any
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

def add_image_with_subtitle(doc: Any, img_stream: Any, caption: str, width_inches: float) -> None:
    """
    Adiciona uma imagem centralizada ao documento Word com uma legenda formatada abaixo.

    A legenda utiliza a fonte Verdana, tamanho 9, na cor preta, com espaçamento de 1.5.

    Args:
        doc (docx.document.Document): O objeto do documento Word (python-docx).
        img_stream (io.BytesIO): O fluxo de dados da imagem ou caminho do arquivo.
        caption (str): O texto da legenda a ser exibido abaixo da imagem.
        width_inches (float): A largura da imagem em polegadas.

    Returns:
        None
    """
    # Adicionando o parágrafo da imagem e centralizando
    p_img = doc.add_paragraph()
    p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_img = p_img.add_run()
    run_img.add_picture(img_stream, width=Inches(width_inches))
    
    # Adicionando o parágrafo da legenda (subtitle)
    p_leg = doc.add_paragraph()
    p_leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_leg.paragraph_format.line_spacing = 1.5
    run_leg = p_leg.add_run(caption)
    run_leg.font.name = 'Verdana'
    run_leg.font.size = Pt(9)
    run_leg.font.color.rgb = RGBColor(0, 0, 0)