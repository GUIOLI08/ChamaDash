import os
import io
import datetime
import base64
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .consult_ia import consult_ia
from .generate_image_graphic import generate_image_graphic
from .add_title import add_title
from .add_paragraph import add_paragraph
from .add_image_with_subtitle import add_image_with_subtitle

def shade_cell(cell, fill_color):
    """Pinta o fundo da célula apagando cores anteriores (Garante o Título visível)"""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn('w:shd'))
    if shd is not None: tcPr.remove(shd)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_borders(cell):
    """Injeta XML bruto para desenhar uma borda preta forte em todas as células"""
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is not None: tcPr.remove(tcBorders)
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        border = OxmlElement(f'w:{edge}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4') # Espessura da borda
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000') # Cor preta absoluta
        tcBorders.append(border)
    tcPr.append(tcBorders)

def format_cell(cell, text, bold=False, color=None, font_size=7):
    """Escreve e formata o texto da tabela"""
    cell.text = str(text)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.font.name = 'Verdana'
    run.font.size = Pt(font_size)
    if bold: run.font.bold = True
    if color: run.font.color.rgb = color

def create_word_report(dados_dashboard):
    caminho_template = "template_netra.docx"
    doc = Document(caminho_template) if os.path.exists(caminho_template) else Document()

    if 'Normal' in doc.styles:
        doc.styles['Normal'].font.name = 'Verdana'
        doc.styles['Normal'].font.size = Pt(9)
        doc.styles['Normal'].paragraph_format.space_after = Pt(6) 
        doc.styles['Normal'].paragraph_format.line_spacing = 1.5 

    for i in range(1, 4):
        style_name = f'Heading {i}'
        if style_name in doc.styles:
            doc.styles[style_name].paragraph_format.space_before = Pt(18)
            doc.styles[style_name].paragraph_format.space_after = Pt(10)

    ai_texts = consult_ia(dados_dashboard)
    current_date = datetime.datetime.now().strftime("%d/%m/%Y")

    for section in doc.sections:
        for p in section.footer.paragraphs:
            if "TIC-RQ-28" in p.text or "Relatório" in p.text:
                p.text = "" 
                run = p.add_run(f"TIC-RQ-28 - Relatório de Nível de Serviço e Indicadores da Qualidade de Serviço {current_date}")
                run.font.name = 'Verdana'
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(100, 100, 100) 
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # --- CAPA ---
    for _ in range(1): doc.add_paragraph()
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_t = p_titulo.add_run("TIC-RQ-28 - Relatório de Nível de Serviço e Indicadores da Qualidade de Serviço")
    run_t.font.name = 'Verdana'
    run_t.font.size = Pt(12)
    run_t.font.bold = True
    
    for _ in range(4): doc.add_paragraph()
    p_cli = doc.add_paragraph()
    p_cli.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_cli = p_cli.add_run("Cliente: Bahiagás")
    run_cli.font.name = 'Verdana'
    run_cli.font.size = Pt(11)
    run_cli.font.bold = True
    
    for _ in range(3): doc.add_paragraph()
    add_paragraph(doc, f"Emitido em: {current_date}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_paragraph(doc, "Período de apuração: [MÊS DE APURAÇÃO]").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- SUMÁRIO ---
    add_title(doc, 'Sumário', 1)
    sumario = [
        "1. INTRODUÇÃO",
        "2. SUMÁRIO DA PERFORMANCE RELACIONADO ÀS METAS DE NÍVEL DE SERVIÇO",
        "3. RESULTADOS APURADOS DOS SERVIÇOS",
        "   3.1 DADOS GERAIS",
        "      3.1.1 Percentual de chamados registrados por tipo",
        "      3.1.2 Percentual de Incidentes e Solicitações de Serviços registradas por prioridade relatada",
        "      3.1.3 Atendimentos por área demandante",
        "4. EVENTOS SIGNIFICATIVOS",
        "   4.1 REGISTROS DE INCIDENTE",
        "      4.1.1 Planos de Continuidade Acionados",
        "      4.1.2 Incidentes Graves",
        "   4.2 REGISTROS DE PROBLEMA",
        "   4.3 CARGA DE TRABALHO",
        "      4.3.1 Serviços mais utilizados",
        "5. ANÁLISE DE TENDÊNCIAS",
        "6. ANEXOS"
    ]
    for item in sumario:
        is_bold = not item.startswith(" ")
        add_paragraph(doc, item, bold=is_bold)
        
    doc.add_page_break()

    # --- 1. INTRODUCTION ---
    add_title(doc, '1. INTRODUÇÃO', 1)
    add_paragraph(doc, ai_texts.get('introducao', 'Lorem ipsum dolor sit amet consectetur adipisicing elit. Consectetur soluta optio nemo omnis, minima similique quasi sit pariatur fuga blanditiis asperiores ab illum non iusto accusantium perferendis officia corporis quibusdam?\n\nLorem ipsum dolor sit amet consectetur adipisicing elit. Consectetur soluta optio nemo omnis, minima similique quasi sit pariatur fuga blanditiis asperiores ab illum non iusto accusantium perferendis officia corporis quibusdam?'))
    doc.add_page_break()
    
    # --- 2. SLA (MATRIZ GERAL) ---
    add_title(doc, '2. SUMÁRIO DA PERFORMANCE RELACIONADO ÀS METAS DE NÍVEL DE SERVIÇO', 1)
    add_paragraph(doc, 'O quadro abaixo resume o nível de serviço obtido no período e compara com o nível proposto:')
    
    t_sla = doc.add_table(rows=7, cols=13)
    t_sla.autofit = False
    t_sla.allow_autofit = False
    
    # Travando as larguras exatas para impedir o Word de quebrar o texto
    larguras = [Inches(1.0)] + [Inches(0.42)] * 12
    for i, col in enumerate(t_sla.columns):
        col.width = larguras[i]
        for cell in col.cells:
            cell.width = larguras[i]
            set_cell_borders(cell) # <-- FORÇA A BORDA PRETA EM TODAS AS CÉLULAS
            
    cor_branca = RGBColor(255, 255, 255)
    cor_preta = RGBColor(0, 0, 0) # <-- GARANTE LEITURA NOS DADOS
    
    # Linha 0 (Título da Tabela)
    t_sla.cell(0,0).merge(t_sla.cell(0,12))
    format_cell(t_sla.cell(0,0), "📊 MATRIZ GERAL DE CHAMADOS E DESEMPENHO (SLA)", bold=True, color=cor_branca, font_size=8)
    shade_cell(t_sla.cell(0,0), "244062")

    # Linha 1 (Grupos)
    headers_r1 = [(0,1,"Chamados"), (2,4,"Baixa"), (5,7,"Média"), (8,10,"Alta"), (11,12,"Total")]
    for c_start, c_end, text in headers_r1:
        c = t_sla.cell(1, c_start)
        c.merge(t_sla.cell(1, c_end))
        format_cell(c, text, bold=True, color=cor_branca, font_size=8)
        shade_cell(c, "244062")

    # Linha 2 (Cabeçalhos)
    headers_r2 = ["Tipos", "Qtd", "No Prazo", "Fora", "ANS", "No Prazo", "Fora", "ANS", "No Prazo", "Fora", "ANS", "No Prazo", "Fora"]
    for idx, h in enumerate(headers_r2):
        format_cell(t_sla.cell(2, idx), h, bold=True, color=cor_branca, font_size=6)
        shade_cell(t_sla.cell(2, idx), "4F81BD")

    matriz_sla = dados_dashboard.get("matriz_sla", {})
    tipos = ["Incidentes", "Solicitações", "Problemas", "Total"]
    
    # Linhas de Dados (Textos FORÇADOS para PRETO)
    for r_idx, tipo in enumerate(tipos, start=3):
        is_bold = (r_idx == 6)
        format_cell(t_sla.cell(r_idx, 0), tipo, bold=is_bold, color=cor_preta, font_size=7)
        valores = matriz_sla.get(tipo, ["-"] * 12)
        
        for c_idx in range(1, 13):
            val = valores[c_idx - 1]
            if val != "-" and c_idx in [4, 7, 10]:
                try: val_str = f"{float(val)*100:.0f}%"
                except: val_str = str(val)
            else:
                val_str = str(val)
            
            # Aqui é onde o texto fica preto!
            format_cell(t_sla.cell(r_idx, c_idx), val_str, bold=is_bold, color=cor_preta, font_size=7)

    p_tab = doc.add_paragraph()
    p_tab.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tab = p_tab.add_run("Tabela 1: Quantitativo de chamados por tipo x impacto")
    run_tab.font.name = 'Verdana'
    run_tab.font.size = Pt(8)

    doc.add_page_break()
    
    # --- 3. RESULTADOS APURADOS ---
    add_title(doc, '3. RESULTADOS APURADOS DOS SERVIÇOS', 1)
    add_title(doc, '3.1 DADOS GERAIS', 2)
    add_paragraph(doc, ai_texts.get('dados_gerais', 'Lorem ipsum dolor sit amet consectetur adipisicing elit. Consectetur soluta optio nemo omnis, minima similique quasi sit pariatur fuga blanditiis asperiores ab illum non iusto accusantium perferendis officia corporis quibusdam?\n\nLorem ipsum dolor sit amet consectetur adipisicing elit. Consectetur soluta optio nemo omnis, minima similique quasi sit pariatur fuga blanditiis asperiores ab illum non iusto accusantium perferendis officia corporis quibusdam?'))
    
    add_title(doc, '3.1.1 Percentual de chamados registrados por tipo', 3)
    doc.paragraphs[-1].paragraph_format.keep_with_next = False 
    
    if dados_dashboard.get('tipos_gerais'):
        img_tipos = generate_image_graphic(dados_dashboard['tipos_gerais'], 'Tipo de Chamado', 'pie')
        add_image_with_subtitle(doc, img_tipos, "Gráfico 1: Percentual por tipo de chamado", 3.2) 
        doc.paragraphs[-2].paragraph_format.line_spacing = 1.0 
        doc.paragraphs[-1].paragraph_format.keep_with_next = False

    doc.add_page_break()

    # --- 3.1.2 GRAFICOS LADO A LADO COM TAMANHO IDENTICO ---
    add_title(doc, '3.1.2 Percentual de Incidentes e Solicitações de Serviços registradas por prioridade relatada', 3)
    doc.paragraphs[-1].paragraph_format.keep_with_next = False 
    
    if dados_dashboard.get('inc_prio') and dados_dashboard.get('req_prio'):
        img_inc = generate_image_graphic(dados_dashboard['inc_prio'], 'Incidentes por Prioridade', 'pie')
        img_req = generate_image_graphic(dados_dashboard['req_prio'], 'Solicitações por Prioridade', 'pie')
        
        tab_graf = doc.add_table(rows=1, cols=2)
        tab_graf.autofit = False
        tab_graf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Travando a tabela pra garantir que ninguém espichará a imagem
        for cell in tab_graf.columns[0].cells: cell.width = Inches(3.2)
        for cell in tab_graf.columns[1].cells: cell.width = Inches(3.2)
        
        # Célula Esquerda
        c_esq = tab_graf.cell(0, 0)
        p_esq = c_esq.paragraphs[0]
        p_esq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_esq = p_esq.add_run()
        r_esq.add_picture(img_inc, width=Inches(3.0)) # <-- Exatamente iguais
        p_sub_esq = c_esq.add_paragraph("Gráfico 2: Incidentes por prioridade")
        p_sub_esq.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub_esq.runs[0].font.name = 'Verdana'
        p_sub_esq.runs[0].font.size = Pt(8)
        
        # Célula Direita
        c_dir = tab_graf.cell(0, 1)
        p_dir = c_dir.paragraphs[0]
        p_dir.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_dir = p_dir.add_run()
        r_dir.add_picture(img_req, width=Inches(3.0)) # <-- Exatamente iguais
        p_sub_dir = c_dir.add_paragraph("Gráfico 3: Solicitações por prioridade")
        p_sub_dir.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_sub_dir.runs[0].font.name = 'Verdana'
        p_sub_dir.runs[0].font.size = Pt(8)
        
        for row in tab_graf.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    p.paragraph_format.line_spacing = 1.0
                    p.paragraph_format.space_after = Pt(0)

    # --- 3.1.3 Chart for Demanding Areas ---
    add_title(doc, '3.1.3 Atendimentos por área demandante', 3)
    doc.paragraphs[-1].paragraph_format.keep_with_next = False
    
    if dados_dashboard.get('top_setores'):
        img_setores = generate_image_graphic(dados_dashboard['top_setores'], 'Atendimentos por Setor - Top 10', 'bar')
        add_image_with_subtitle(doc, img_setores, "Gráfico 4: Atendimentos abertos por áreas demandantes", 4.5)
        doc.paragraphs[-2].paragraph_format.line_spacing = 1.0
        doc.paragraphs[-1].paragraph_format.keep_with_next = False 
    
    doc.add_page_break()
    
    # --- 4. EVENTOS SIGNIFICATIVOS ---
    add_title(doc, '4. EVENTOS SIGNIFICATIVOS', 1)
    add_title(doc, '4.1 REGISTRO DE INCIDENTE', 2)
    add_title(doc, '4.1.1 Planos de Continuidade Acionados', 3)
    add_paragraph(doc, 'Nesse período não houve acionamento de Plano de Continuidade.')
    
    add_title(doc, '4.1.2 Incidentes Graves', 3)
    add_paragraph(doc, 'Nesse período não houve registro de incidentes graves.')
    
    add_title(doc, '4.2 REGISTRO DE PROBLEMA', 2)
    add_paragraph(doc, 'Nesse período não houve registro de Problemas.')
    
    add_title(doc, '4.3 CARGA DE TRABALHO', 2)
    doc.paragraphs[-1].paragraph_format.keep_with_next = True 
    
    add_title(doc, '4.3.1 Serviços mais utilizados', 3)
    doc.paragraphs[-1].paragraph_format.keep_with_next = True 
    
    p_carga = add_paragraph(doc, 'Houve maior incidência das seguintes Solicitações de Serviço:')
    p_carga.paragraph_format.keep_with_next = True 
    
    if dados_dashboard.get('top_categorias'):
        img_cat = generate_image_graphic(dados_dashboard['top_categorias'], 'Chamados por Categoria', 'bar')
        add_image_with_subtitle(doc, img_cat, "Gráfico 5: Chamados abertos por categoria", 4.5)

    # --- 5. TREND ANALYSIS ---
    add_title(doc, '5. ANÁLISE DE TENDÊNCIAS', 1)
    p_tendencia = doc.add_paragraph()
    run_tend = p_tendencia.add_run(ai_texts.get('dados_gerais', 'Lorem ipsum dolor sit amet consectetur adipisicing elit. Consectetur soluta optio nemo omnis, minima similique quasi sit pariatur fuga blanditiis asperiores ab illum non iusto accusantium perferendis officia corporis quibusdam?\n\nLorem ipsum dolor sit amet consectetur adipisicing elit. Consectetur soluta optio nemo omnis, minima similique quasi sit pariatur fuga blanditiis asperiores ab illum non iusto accusantium perferendis officia corporis quibusdam?'))
    run_tend.font.name = 'Verdana'
    run_tend.font.size = Pt(9)

    # --- 6. ATTACHMENTS ---
    add_title(doc, 'Anexos', 1)
    add_paragraph(doc, 'Insira aqui evidências ou imagens complementares da Central Telefônica ou outros sistemas.')

    word_output = io.BytesIO()
    doc.save(word_output)
    word_output.seek(0)
    return base64.b64encode(word_output.getvalue()).decode("utf-8")