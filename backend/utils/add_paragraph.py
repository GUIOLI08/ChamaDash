from typing import Any
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

def add_paragraph(doc: Any, text: str, bold: bool = False) -> Any:
    """
    Adiciona um parágrafo ao documento utilizando o estilo padrão do projeto.

    O parágrafo utiliza a fonte Verdana, tamanho 9 e cor preta.

    Args:
        doc (docx.document.Document): O objeto do documento Word (python-docx).
        text (str): O conteúdo de texto do parágrafo.
        bold (bool, opcional): Se o texto deve ser em negrito. O padrão é False.

    Returns:
        docx.text.paragraph.Paragraph: O objeto do parágrafo criado.
    """
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Verdana'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0, 0, 0)
    if bold:
        run.font.bold = True
    return p